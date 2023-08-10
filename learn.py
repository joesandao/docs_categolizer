from transformers import BertModel, BertJapaneseTokenizer
import os
import glob
import torch
from docx import Document
from pdfminer.high_level import extract_text

# 特定のパスからすべてのフォルダを取得
base_path = 'rowdata'  # 自分の環境に合わせて変更してください
folders = [os.path.join(base_path, name) for name in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, name))]

label_mapping = {folder: i for i, folder in enumerate(folders)}  # ラベルを数値にマッピング
data = []

tokenizer = BertJapaneseTokenizer.from_pretrained('cl-tohoku/bert-base-japanese-whole-word-masking')
model = BertModel.from_pretrained('cl-tohoku/bert-base-japanese-whole-word-masking')

for folder in folders:
    for file_path in glob.glob(os.path.join(folder, '*')):
        # ファイル拡張子に応じてテキストを抽出
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
        elif file_path.endswith('.pdf'):
            text = extract_text(file_path)

        # テキストをBERTでベクトル化
        inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True, max_length=512)
        with torch.no_grad():
            outputs = model(**inputs)
        
        # [CLS]トークンを抽出（文書全体を表現）
        cls_output = outputs[0][0][0]
        data.append((cls_output.numpy(), label_mapping[folder]))  # ベクトルとラベルの組み合わせを保存

from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score
import numpy as np

# データを特徴ベクトルとラベルに分割
features, labels = zip(*data)
features = np.stack(features)
labels = np.array(labels)

# データを訓練セットとテストセットに分割
X_train, X_test, y_train, y_test = train_test_split(features, labels, test_size=0.2, random_state=42)

# モデルの訓練
clf = LogisticRegression(random_state=0).fit(X_train, y_train)

# テストデータでの予測
y_pred = clf.predict(X_test)

# 予測の精度を計算
accuracy = accuracy_score(y_test, y_pred)
print(f'Accuracy: {accuracy}')

# モデルの保存
from joblib import dump
dump(clf, 'model.joblib') 



import shutil

# ラベルの逆マッピングを作成
reverse_label_mapping = {v: k for k, v in label_mapping.items()}

# 新しい文書のパス
new_doc_path = 'target/情報セキュリティ応用2023#13_爰川.pdf'

# 新しい文書のテキストを抽出
if new_doc_path.endswith('.docx'):
    doc = Document(new_doc_path)
    text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
elif new_doc_path.endswith('.pdf'):
    text = extract_text(new_doc_path)

# テキストをBERTでベクトル化
inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True, max_length=512)
with torch.no_grad():
    outputs = model(**inputs)

# [CLS]トークンを抽出（文書全体を表現）
cls_output = outputs[0][0][0]

# モデルを使ってラベルを予測
predicted_label = clf.predict(cls_output.numpy().reshape(1, -1))

# 分類先のディレクトリ名を取得
output_folder = reverse_label_mapping[predicted_label[0]]

# 分類先のディレクトリを作成（すでに存在する場合は何もしない）
os.makedirs(os.path.join('output', output_folder), exist_ok=True)

# 元の文書を分類先のディレクトリにコピー
shutil.copy(new_doc_path, os.path.join('output', output_folder))

